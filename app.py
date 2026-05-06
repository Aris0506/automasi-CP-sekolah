import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import json
from PyPDF2 import PdfReader
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

# Konfigurasi Halaman
st.set_page_config(page_title="Generator Modul Ajar", page_icon="✈️", layout="centered")

# --- UI BAGIAN SAMPING (API KEY) ---
with st.sidebar:
    st.header("⚙️ Pengaturan Sistem")
    api_key = st.text_input("Masukkan Gemini API Key Anda", type="password", help="API Key gratis tidak akan kami simpan.")
    st.markdown("[Dapatkan Gemini API Key Gratis di sini](https://aistudio.google.com/app/apikey)")
    st.markdown("---")
    st.caption("Aplikasi ini menggunakan mode BYOK (Bring Your Own Key) agar tetap gratis selamanya.")

# --- FUNGSI BACA PDF ---
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

# --- FUNGSI HELPER FORMATTING (WIN-WIN SOLUTION) ---
def tulis_rapi(container, data, rata_kanan_kiri=False):
    # Cek apakah data dari AI berupa List atau String tunggal
    if isinstance(data, list):
        for baris in data:
            bersih = str(baris).strip().replace('**', '')
            if bersih:
                p = container.add_paragraph(bersih)
                if rata_kanan_kiri:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        baris_baris = str(data).split('\n')
        for baris in baris_baris:
            bersih = baris.strip().replace('**', '')
            if bersih:
                p = container.add_paragraph(bersih)
                if rata_kanan_kiri:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# --- FUNGSI RAKIT WORD ---
def create_word_document(mapel, topik, fase, alokasi_waktu, nama_sekolah, nama_guru, nip_guru, nama_kepsek, nip_kepsek, tujuan, langkah, asesmen):
    doc = Document()
    
    # --- IDENTITAS UMUM (Rata Kiri Default) ---
    doc.add_heading("Informasi Umum", level=1)
    doc.add_paragraph(f"Nama Sekolah\t\t: {nama_sekolah}")
    doc.add_paragraph(f"Mata Pelajaran\t\t: {mapel}")
    doc.add_paragraph(f"Fase / Kelas\t\t: {fase}")
    doc.add_paragraph(f"Topik Materi\t\t: {topik}")
    doc.add_paragraph(f"Alokasi Waktu\t\t: {alokasi_waktu}")
    
    # --- 1. TUJUAN PEMBELAJARAN (Justify) ---
    doc.add_heading("A. Tujuan Pembelajaran", level=1)
    tulis_rapi(doc, tujuan, True)
    
    # --- 2. LANGKAH PEMBELAJARAN (Justify dalam Tabel) ---
    doc.add_heading("B. Langkah Pembelajaran", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Tahap Pembelajaran'
    hdr_cells[1].text = 'Kegiatan Inti'
    
    row_cells = table.add_row().cells
    row_cells[0].text = 'Inti'
    
    cell_kegiatan = row_cells[1]
    cell_kegiatan.text = "" 
    tulis_rapi(cell_kegiatan, langkah, True)
    
    # --- 3. ASESMEN (Justify) ---
    doc.add_heading("C. Asesmen", level=1)
    tulis_rapi(doc, asesmen, True)
    
    # --- 4. BLOK TANDA TANGAN (VERSI PRESISI SIMETRIS) ---
    doc.add_paragraph("\n") 
    ttd_table = doc.add_table(rows=1, cols=2)
    
    # Bagian Kiri (Kepsek)
    cell_kiri = ttd_table.cell(0, 0)
    p_kiri = cell_kiri.paragraphs[0]
    p_kiri.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Total ada 2 baris teks di awal (Mengetahui + Kepala Sekolah)
    p_kiri.add_run("Mengetahui,\n")
    p_kiri.add_run("Kepala Sekolah\n\n\n\n") # Spasi TTD
    run_kps = p_kiri.add_run(nama_kepsek)
    run_kps.bold = True
    p_kiri.add_run(f"\nNIP. {nip_kepsek}")
    
    # Bagian Kanan (Guru)
    cell_kanan = ttd_table.cell(0, 1)
    p_kanan = cell_kanan.paragraphs[0]
    p_kanan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # KUNCINYA DI SINI: Tambah \n kosong di awal biar sejajar sama "Mengetahui,"
    p_kanan.add_run("\n") 
    p_kanan.add_run("Guru Mata Pelajaran\n\n\n\n") # Spasi TTD
    run_guru = p_kanan.add_run(nama_guru)
    run_guru.bold = True
    p_kanan.add_run(f"\nNIP. {nip_guru}")

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# --- UI BAGIAN UTAMA ---
st.title("Generator Modul Ajar (RAG Powered)")

st.subheader("📁 Dokumen Sumber (CP)")
uploaded_file = st.file_uploader("Upload PDF Capaian Pembelajaran (CP)", type="pdf")

if uploaded_file:
    st.success(f"File '{uploaded_file.name}' berhasil diunggah.")

st.subheader("🏫 Identitas Sekolah & Guru")
col_id1, col_id2 = st.columns(2)
with col_id1:
    nama_sekolah = st.text_input("Nama Sekolah", placeholder="Contoh: SMA Negeri 1")
    nama_kepsek = st.text_input("Nama Kepala Sekolah", placeholder="Nama Lengkap & Gelar")
    nip_kepsek = st.text_input("NIP Kepala Sekolah")
with col_id2:
    nama_guru = st.text_input("Nama Guru Mata Pelajaran")
    nip_guru = st.text_input("NIP Guru")
    alokasi_waktu = st.text_input("Alokasi Waktu", placeholder="Contoh: 2 x 45 Menit")

st.subheader("📝 Detail Modul")
col1, col2 = st.columns(2)
with col1:
    mapel = st.text_input("Mata Pelajaran", placeholder="Contoh: Informatika")
    fase = st.selectbox("Fase", ["A", "B", "C", "D", "E", "F"])
with col2:
    topik = st.text_input("Topik Materi", placeholder="Contoh: Berpikir Komputasional")

# --- LOGIKA TOMBOL GENERATE ---
if st.button("Generate Modul Ajar", type="primary", use_container_width=True):
    if not api_key or not uploaded_file:
        st.error("Pastikan API Key dan File PDF CP sudah diisi/diunggah!")
    elif not mapel or not topik:
        st.warning("Harap isi Mata Pelajaran dan Topik Materi.")
    else:
        with st.spinner("Sedang membedah PDF dan merakit modul..."):
            try:
                context_cp = extract_text_from_pdf(uploaded_file)
                
                genai.configure(api_key=api_key)
                # Menggunakan model versi terbaru sesuai pengecekan lo tadi
                model = genai.GenerativeModel('gemini-2.5-flash')
                
                prompt = f"""
                Anda adalah pakar kurikulum Merdeka.
                Gunakan data Capaian Pembelajaran (CP) berikut:
                ---
                Konteks CP: {context_cp[:5000]}
                ---
                Buatkan Modul Ajar untuk:
                Mata Pelajaran: {mapel}
                Fase: {fase}
                Topik: {topik}
                
                ATURAN FORMAT OUTPUT:
                1. DILARANG KERAS menggunakan Markdown (JANGAN gunakan ** untuk huruf tebal, JANGAN gunakan - atau * untuk bullet points).
                2. Gunakan teks biasa murni.
                3. Jika butuh membuat daftar poin, gunakan penomoran angka biasa (1. 2. 3. dan seterusnya).
                4. Pisahkan SETIAP POIN materi dengan ENTER (baris baru). Jangan digabung dalam satu paragraf!
                
                Output HARUS dalam format JSON murni tanpa markdown, dengan struktur key:
                "tujuan", "langkah", "asesmen"
                """
                
                # --- KONFIGURASI ENGINE (LEBIH STABIL & DETERMINISTIK) ---
                config = genai.GenerationConfig(
                    temperature=0.0,            # Biar gak halu dan tetep kaku (akurat)[cite: 1, 6]
                    top_p=0.95,
                    top_k=40,
                    max_output_tokens=4096,     # Jatah biar penjelasan gak kepotong di tengah
                    response_mime_type="application/json" # Tetap wajib JSON
                )

                # Panggil model pake config yang udah lengkap tadi
                response = model.generate_content(
                    prompt,
                    generation_config=config
                )
                
                ai_data = json.loads(response.text)
                
                tujuan_ai = ai_data.get("tujuan", "Gagal memproses tujuan.")
                langkah_ai = ai_data.get("langkah", "Gagal memproses langkah.")
                asesmen_ai = ai_data.get("asesmen", "Gagal memproses asesmen.")
                
                docx_file = create_word_document(
                    mapel, topik, fase, alokasi_waktu, 
                    nama_sekolah, nama_guru, nip_guru, nama_kepsek, nip_kepsek, 
                    tujuan_ai, langkah_ai, asesmen_ai
                )
                
                st.success("✅ Modul Berhasil Dibuat!")
                
                st.download_button(
                    label="⬇️ Download Modul (.docx)",
                    data=docx_file,
                    file_name=f"Modul_{mapel}_{topik}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                
            except json.JSONDecodeError:
                st.error("AI gagal menghasilkan format JSON yang benar. Silakan coba klik generate lagi.")
            except Exception as e:
                st.error(f"Terjadi kesalahan sistem: {e}")